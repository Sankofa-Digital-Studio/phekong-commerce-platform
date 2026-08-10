from pathlib import Path
from PIL import Image, ImageOps, ImageEnhance, ImageFilter, ImageDraw
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
BASE = ROOT / "docs" / "media-scan"
SRC = BASE / "source"
OUT = BASE / "edited"
OUT.mkdir(parents=True, exist_ok=True)

def prep(name, size, bg, panel, product_box):
    im = ImageOps.exif_transpose(Image.open(SRC / name).convert("RGB"))
    im = ImageEnhance.Contrast(im).enhance(1.05)
    im = ImageEnhance.Color(im).enhance(0.96)
    canvas = Image.new("RGB", size, bg)
    d = ImageDraw.Draw(canvas)
    x, y, w, h = product_box
    shadow = Image.new("RGBA", size, (0,0,0,0)); sd=ImageDraw.Draw(shadow)
    sd.rounded_rectangle((x+16,y+18,x+w+16,y+h+18), 34, fill=(38,30,20,38))
    canvas = Image.alpha_composite(canvas.convert("RGBA"), shadow.filter(ImageFilter.GaussianBlur(18)))
    d = ImageDraw.Draw(canvas); d.rounded_rectangle((x,y,x+w,y+h),34,fill=panel)
    fitted = ImageOps.contain(im, (w-56,h-56), Image.Resampling.LANCZOS)
    canvas.alpha_composite(fitted.convert("RGBA"), (x+(w-fitted.width)//2,y+(h-fitted.height)//2))
    return canvas.convert("RGB")

assets = [
    ("IMG_0077.JPG", "home-collection-hero-aloe-puree.jpg", (1600,900), "#E8D8C2", "#F7F2E9", (900,70,560,760)),
    ("IMG_0081.JPG", "product-detail-aloe-puree.jpg", (1200,1200), "#D7DDCF", "#F5F1E8", (170,90,860,1020)),
    ("IMG_0068.JPG", "catalogue-card-herbal-therapy.jpg", (900,1125), "#C9D1B7", "#F7F3EA", (110,95,680,900)),
]
for src, dest, size, bg, panel, box in assets:
    prep(src,size,bg,panel,box).save(OUT/dest, quality=91, optimize=True)

md = """# Frontend media scan and usage plan

**Status:** local review asset pack; not connected to the live frontend.
**Source:** shared Google Drive folder supplied by the product owner.
**Branch:** `sankofa_xciv/feature-media-scan-frontend`

## Decision summary

Three JPEG frames were selected from the paired JPEG/CR2 shoot. The edits are deterministic: EXIF orientation correction, modest contrast/colour balancing, responsive crops, and presentation on warm/sage studio panels. Product pixels and label artwork remain unchanged. An AI background concept was tested but rejected because it distorted small label text.

## Screen assignments

| Screen | Edited asset | Source | Role | Recommended treatment |
|---|---|---|---|---|
| Homepage / collection landing | `edited/home-collection-hero-aloe-puree.jpg` | `IMG_0077.JPG` | Primary campaign hero | 16:9; bottle right; left side reserved for copy and CTA |
| Product detail | `edited/product-detail-aloe-puree.jpg` | `IMG_0081.JPG` | Main gallery image | Square; neutral sage surround; use `object-fit: contain` |
| Product listing / editorial tile | `edited/catalogue-card-herbal-therapy.jpg` | `IMG_0068.JPG` | Product card or category story | 4:5; strong mobile crop; label remains central |

## Frontend usage contract

- Keep source masters outside `public/`; only approved derivatives should be promoted later.
- Before production use, confirm product names, sizes, claims, ownership/consent, and alt text with the catalogue owner.
- Proposed alt text: “Large bottle of Aloe Puree complementary supplement”; “Compact bottle of Aloe Puree complementary supplement”; “Bottle of Herbal Therapy herbal tea”.
- Serve responsive AVIF/WebP variants at roughly 480, 768, 1200, and 1600 px; retain JPEG as fallback.
- Do not place text over labels. Hero copy belongs in the reserved negative-space area.

## Three-pass critique

1. **Selection pass:** chose frames with the clearest silhouette, legible hero product name, and distinct screen utility; rejected near-duplicates.
2. **Usability pass:** corrected orientation and created aspect ratios for desktop and mobile without altering label pixels.
3. **Risk pass:** AI recreation was rejected for label fidelity; production promotion remains gated by product-data and rights confirmation.

## Not selected

`IMG_0050.JPG`, `IMG_0039.JPG`, and `IMG_0034.JPG` were retained only as scan references because the local downloads showed lower reliability or duplicated the selected visual roles.
"""
(BASE/"frontend-media-scan.md").write_text(md, encoding="utf-8")

cards=''.join(f'<article><img src="edited/{dest}" alt="{alt}"><h2>{title}</h2><p><b>Source:</b> {src}</p><p>{role}</p></article>' for src,dest,_,_,_,_ ,title,role,alt in [
    (*assets[0],"Homepage / collection hero","16:9 composition with copy space on the left.","Large bottle of Aloe Puree complementary supplement"),
    (*assets[1],"Product detail gallery","Square presentation for a primary product-detail image.","Compact bottle of Aloe Puree complementary supplement"),
    (*assets[2],"Catalogue card / editorial tile","Mobile-first 4:5 crop for product discovery.","Bottle of Herbal Therapy herbal tea")])
html=f'''<!doctype html><html lang="en"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1,viewport-fit=cover"><title>Phekong frontend media scan</title><style>:root{{--ink:#27251f;--paper:#f6f1e8;--sage:#65725a}}*{{box-sizing:border-box}}body{{margin:0;font:16px/1.55 system-ui,-apple-system,BlinkMacSystemFont,"Segoe UI",sans-serif;background:var(--paper);color:var(--ink)}}header{{padding:clamp(2rem,7vw,6rem);background:#263127;color:white}}main{{max-width:1180px;margin:auto;padding:clamp(1rem,4vw,3rem)}}.grid{{display:grid;grid-template-columns:repeat(auto-fit,minmax(min(100%,280px),1fr));gap:1.25rem}}article{{background:white;border-radius:20px;overflow:hidden;box-shadow:0 10px 30px #30251418}}img{{width:100%;aspect-ratio:4/3;object-fit:cover;display:block}}h1{{max-width:18ch;font-size:clamp(2rem,6vw,4.5rem);line-height:1.05}}h2,p{{margin:1rem 1.25rem}}.note{{border-left:4px solid var(--sage);padding:1rem;background:#fff}}@media(max-width:600px){{article{{border-radius:14px}}img{{aspect-ratio:4/5}}}}</style></head><body><header><p>Local-only media selection pack</p><h1>Product photography mapped to frontend screens</h1></header><main><p class="note">No assets in this pack are connected to the live site. Product and rights approval is required before promotion.</p><section class="grid">{cards}</section></main></body></html>'''
(BASE/"frontend-media-scan.html").write_text(html,encoding="utf-8")

doc=Document(); sec=doc.sections[0]; sec.top_margin=Inches(.7); sec.bottom_margin=Inches(.7)
styles=doc.styles; styles['Normal'].font.name='Aptos'; styles['Normal'].font.size=Pt(10.5)
for s,c in [('Title','263127'),('Heading 1','65725A'),('Heading 2','65725A')]: styles[s].font.color.rgb=RGBColor.from_string(c)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('PHEKONG'); r.bold=True; r.font.size=Pt(14); r.font.color.rgb=RGBColor.from_string('65725A')
doc.add_heading('Frontend Media Scan & Usage Plan',0); doc.add_paragraph('Local-only product image selection pack | 9 August 2026').alignment=WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break(); doc.add_heading('Table of contents',1); p=doc.add_paragraph(); run=p.add_run(); fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'TOC \\o "1-3" \\h \\z \\u'); run._r.addnext(fld)
doc.add_page_break(); doc.add_heading('1. Decision summary',1); doc.add_paragraph('Three JPEG frames were selected from paired JPEG/CR2 captures. Orientation, colour, crop, and presentation panels were adjusted without repainting product labels. An AI concept was rejected because small label text changed.')
for i,(src,dest,_,_,_,_) in enumerate(assets,1):
    title=['Homepage / collection landing','Product-detail gallery','Catalogue card / editorial tile'][i-1]
    doc.add_heading(f'{i+1}. {title}',1); doc.add_picture(str(OUT/dest),width=Inches(6.2)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Source: {src} | Edited asset: {dest}')
    doc.add_paragraph(['Use the 16:9 composition as a campaign hero, keeping marketing copy in the left negative space.','Use the square composition as the main gallery image with contain behavior and a zoomable original.','Use the 4:5 crop for product grids and editorial discovery modules, especially on mobile.'][i-1])
doc.add_heading('5. Production gate',1); doc.add_paragraph('Before frontend promotion, verify catalogue mapping, product claims, image ownership/consent, final alt text, and responsive AVIF/WebP exports. This branch intentionally changes no application code.')
doc.add_heading('6. Three-pass quality review',1)
for t in ['Selection: distinct utility and readable product identity.','Usability: desktop/mobile aspect ratios and copy-safe placement.','Risk: rejected AI label distortion; production remains approval-gated.']: doc.add_paragraph(t,style='List Bullet')
doc.save(BASE/'Phekong_Frontend_Media_Scan_and_Usage_Plan.docx')
print('generated')
